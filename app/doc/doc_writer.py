import threading
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from app.event_bus import EventBus
from app.transcription.timestamp_utils import seconds_to_hms, format_frame_tag, video_session_id

DEFAULT_OUTPUT = Path("output") / f"notes_{video_session_id()}.docx"
FONT_NAME = "Calibri"
FONT_SIZE_BODY = 11
FONT_SIZE_HEADING = 13


class DocWriter:
    """Manages a live .docx document with thread-safe writes.

    All mutation methods acquire self._lock. Call save() or rely on auto-save
    (triggered by a QTimer in FloatingWidget every 60 seconds).
    """

    def __init__(self, event_bus: EventBus, output_path: Path | None = None):
        self._bus = event_bus
        self._path = output_path or DEFAULT_OUTPUT
        self._path.parent.mkdir(parents=True, exist_ok=True)
        self._doc = Document()
        self._lock = threading.Lock()
        self._setup_styles()

        # Auto-wire incoming signals
        self._bus.transcript_final.connect(self._on_transcript_final)
        self._bus.ocr_result_ready.connect(self._on_ocr_result)
        self._bus.summary_ready.connect(self.ai_summary)

    # -----------------------------------------------------------------
    # Document setup
    # -----------------------------------------------------------------

    def _setup_styles(self) -> None:
        style = self._doc.styles["Normal"]
        style.font.name = FONT_NAME
        style.font.size = Pt(FONT_SIZE_BODY)

    # -----------------------------------------------------------------
    # Append operations
    # -----------------------------------------------------------------

    def append_transcript(self, text: str, timestamp: float) -> None:
        with self._lock:
            tag = format_frame_tag(timestamp)
            p = self._doc.add_paragraph()
            run = p.add_run(f"{tag} ")
            run.bold = True
            run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
            p.add_run(text)
            self._style_paragraph(p, FONT_SIZE_BODY)

    def append_note(self, text: str) -> None:
        with self._lock:
            p = self._doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            self._style_paragraph(p, FONT_SIZE_HEADING)

    def append_ocr(
        self, text: str, image_path: str, timestamp: float
    ) -> None:
        with self._lock:
            # Section heading
            heading = self._doc.add_paragraph()
            run = heading.add_run(f"[OCR @ {seconds_to_hms(timestamp)}]")
            run.bold = True
            run.font.color.rgb = RGBColor(0x20, 0x60, 0xA0)

            # Image inline if file exists and is reasonably sized
            try:
                import os
                if os.path.exists(image_path) and os.path.getsize(image_path) < 5_000_000:
                    self._doc.add_picture(image_path, width=Inches(4))
            except Exception:
                pass

            # OCR text block
            p = self._doc.add_paragraph(text)
            self._style_paragraph(p, FONT_SIZE_BODY)

    # -----------------------------------------------------------------
    # Insert at position
    # -----------------------------------------------------------------

    def insert_at_line(self, line_number: int, text: str) -> None:
        """Insert a paragraph before the given 1-based paragraph index."""
        with self._lock:
            paragraphs = self._doc.paragraphs
            idx = max(0, min(line_number - 1, len(paragraphs)))
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = text
            new_r.append(new_t)
            new_p.append(new_r)
            if idx < len(paragraphs):
                paragraphs[idx]._p.addprevious(new_p)
            else:
                self._doc.element.body.append(new_p)

    def insert_at_page(self, page_number: int, text: str) -> None:
        """Insert text before the Nth page break (1-based).

        Note: OOXML has no rendered page boundary concept. This counts
        explicit w:br[@w:type='page'] elements as page separators.
        """
        with self._lock:
            body = self._doc.element.body
            ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            breaks = body.findall(".//w:br[@w:type='page']", ns)

            target_idx = page_number - 1
            if target_idx < len(breaks):
                anchor = breaks[target_idx].getparent()
            else:
                anchor = None

            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = text
            new_r.append(new_t)
            new_p.append(new_r)

            if anchor is not None:
                anchor.addprevious(new_p)
            else:
                body.append(new_p)

    # -----------------------------------------------------------------
    # Reformatting
    # -----------------------------------------------------------------

    def reformat(self) -> None:
        """Normalize fonts, collapse duplicate blank lines, align paragraphs."""
        with self._lock:
            prev_empty = False
            to_remove = []
            for p in self._doc.paragraphs:
                is_empty = not p.text.strip()
                if is_empty and prev_empty:
                    to_remove.append(p._element)
                prev_empty = is_empty
                self._style_paragraph(p, FONT_SIZE_BODY)

            for el in to_remove:
                el.getparent().remove(el)

            self._bus.status_changed.emit("Document reformatted.")

    # -----------------------------------------------------------------
    # AI Summary
    # -----------------------------------------------------------------

    def ai_summary(self, summary_text: str) -> None:
        """Insert an AI summary section at the beginning of the document."""
        with self._lock:
            body = self._doc.element.body
            first_child = body[0] if len(body) else None

            # Heading paragraph
            heading_p = OxmlElement("w:p")
            heading_r = OxmlElement("w:r")
            heading_rpr = OxmlElement("w:rPr")
            heading_b = OxmlElement("w:b")
            heading_sz = OxmlElement("w:sz")
            heading_sz.set(qn("w:val"), str(FONT_SIZE_HEADING * 2))
            heading_rpr.append(heading_b)
            heading_rpr.append(heading_sz)
            heading_t = OxmlElement("w:t")
            heading_t.text = "AI Summary"
            heading_r.append(heading_rpr)
            heading_r.append(heading_t)
            heading_p.append(heading_r)

            # Summary text paragraph
            summary_p = OxmlElement("w:p")
            summary_r = OxmlElement("w:r")
            summary_t = OxmlElement("w:t")
            summary_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            summary_t.text = summary_text
            summary_r.append(summary_t)
            summary_p.append(summary_r)

            # Separator page break after summary
            sep_p = OxmlElement("w:p")
            sep_r = OxmlElement("w:r")
            sep_br = OxmlElement("w:br")
            sep_br.set(qn("w:type"), "page")
            sep_r.append(sep_br)
            sep_p.append(sep_r)

            if first_child is not None:
                first_child.addprevious(heading_p)
                heading_p.addnext(summary_p)
                summary_p.addnext(sep_p)
            else:
                body.append(heading_p)
                body.append(summary_p)
                body.append(sep_p)

        self._bus.status_changed.emit("AI summary inserted.")

    # -----------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------

    def save(self, path: Path | None = None) -> None:
        target = path or self._path
        with self._lock:
            self._doc.save(str(target))
        self._bus.doc_saved.emit(str(target))
        self._bus.status_changed.emit(f"Saved: {target.name}")

    def get_path(self) -> Path:
        return self._path

    # -----------------------------------------------------------------
    # Signal handlers
    # -----------------------------------------------------------------

    def _on_transcript_final(self, text: str) -> None:
        self.append_transcript(text, 0.0)

    def _on_ocr_result(self, text: str, image_path: str) -> None:
        import time
        self.append_ocr(text, image_path, time.monotonic())

    # -----------------------------------------------------------------
    # Internal helpers
    # -----------------------------------------------------------------

    def _style_paragraph(self, p, font_size: int) -> None:
        for run in p.runs:
            run.font.name = FONT_NAME
            run.font.size = Pt(font_size)
